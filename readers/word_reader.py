from io import BytesIO
import tempfile

from docx import Document

try:
    import textract
except ImportError:
    textract = None

from .base import BaseReader


class WordReader(BaseReader):

    extensions = ["doc", "docx"]

    def read(self, filename, contents):

        ext = filename.split(".")[-1].lower()

        if ext == "docx":
            document = Document(BytesIO(contents))

            text = []

            for para in document.paragraphs:
                if para.text.strip():
                    text.append(para.text)

            for table in document.tables:
                for row in table.rows:
                    text.append(
                        " | ".join(cell.text for cell in row.cells)
                    )

            return "\n".join(text)

        if ext == "doc":

            if textract is None:
                return ".doc support not installed."

            with tempfile.NamedTemporaryFile(suffix=".doc") as f:
                f.write(contents)
                f.flush()

                return textract.process(
                    f.name
                ).decode(errors="ignore")